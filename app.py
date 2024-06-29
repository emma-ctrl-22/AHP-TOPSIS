from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
import numpy as np
from docx import Document
import os
import io
import logging
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# Configure CORS
origins = ["http://localhost:3000"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def calculate_eigenvector(matrix):
    eigenvalues, eigenvectors = np.linalg.eig(matrix)
    max_eigenvalue = np.max(eigenvalues).real
    max_eigenvector = eigenvectors[:, np.argmax(eigenvalues)].real
    normalized_vector = max_eigenvector / np.sum(max_eigenvector)
    return max_eigenvalue, normalized_vector

def calculate_consistency(matrix, max_eigenvalue):
    n = matrix.shape[0]
    ci = (max_eigenvalue - n) / (n - 1)
    ri_values = {1: 0.00, 2: 0.00, 3: 0.58, 4: 0.90, 5: 1.12, 6: 1.24, 7: 1.32, 8: 1.41, 9: 1.45, 10: 1.49}
    ri = ri_values.get(n, 1.49)
    cr = ci / ri if ri != 0 else 0
    return ci, cr, ri

def geometric_mean(matrices):
    product_matrix = np.ones_like(matrices[0])
    for matrix in matrices:
        product_matrix *= matrix
    return np.power(product_matrix, 1/len(matrices))

def create_matrices(file_content):
    logger.info("Creating matrices from file content")
    data = pd.read_csv(io.StringIO(file_content), header=None)
    numeric_data = data.apply(pd.to_numeric, errors='coerce').dropna()

    num_engineers = numeric_data.shape[0]
    engineer_matrices = {}

    for i in range(num_engineers):
        sub_criteria = numeric_data.iloc[i].values
        num_sub_criteria = int((1 + (1 + 8 * len(sub_criteria))**0.5) / 2)

        if not (num_sub_criteria * (num_sub_criteria - 1) // 2) == len(sub_criteria):
            logger.error(f"Invalid number of sub-criteria for engineer {i + 1}. The data does not form a valid matrix.")
            raise ValueError(f"Invalid number of sub-criteria for engineer {i + 1}. The data does not form a valid matrix.")

        matrix = np.eye(num_sub_criteria)
        k = 0
        for row in range(num_sub_criteria):
            for col in range(row + 1, num_sub_criteria):
                if k >= len(sub_criteria):
                    raise IndexError(f"Index {k} is out of bounds for sub_criteria with length {len(sub_criteria)}")
                matrix[row, col] = sub_criteria[k]
                matrix[col, row] = 1 / sub_criteria[k]
                k += 1
        engineer_matrices[f'Engineer {i + 1}'] = matrix

    return engineer_matrices

def process_criteria(file_content):
    logger.info("Processing criteria")
    engineer_matrices = create_matrices(file_content)
    consistent_matrices = []
    results = []

    for engineer, matrix in engineer_matrices.items():
        max_eigenvalue, weights = calculate_eigenvector(matrix)
        ci, cr, ri = calculate_consistency(matrix, max_eigenvalue)
        if cr <= 0.1:
            consistent_matrices.append(matrix)
        results.append({
            'engineer': engineer,
            'matrix': matrix.round(3).tolist(),
            'weights': weights.round(3).tolist(),
            'max_eigenvalue': round(max_eigenvalue, 3),
            'ci': round(ci, 3),
            'cr': round(cr, 3),
            'ri': ri
        })

    if consistent_matrices:
        aggregate_matrix = geometric_mean(consistent_matrices)
        max_eigenvalue, weights = calculate_eigenvector(aggregate_matrix)
        ci, cr, ri = calculate_consistency(aggregate_matrix, max_eigenvalue)
        aggregate_result = {
            'aggregate_matrix': aggregate_matrix.round(3).tolist(),
            'weights': weights.round(3).tolist(),
            'max_eigenvalue': round(max_eigenvalue, 3),
            'ci': round(ci, 3),
            'cr': round(cr, 3),
            'ri': ri
        }
    else:
        aggregate_result = None

    return results, aggregate_result

def save_to_word(results, aggregate_result):
    logger.info("Saving results to Word document")
    doc = Document()
    doc.add_heading('Criteria Analysis', level=1)

    for result in results:
        doc.add_heading(result['engineer'], level=2)
        doc.add_paragraph(f"Pairwise Comparison Matrix:\n{np.array(result['matrix'])}")
        doc.add_paragraph(f"Weights:\n{result['weights']}")
        doc.add_paragraph(f"Max Eigenvalue: {result['max_eigenvalue']}")
        doc.add_paragraph(f"Consistency Index (CI): {result['ci']}")
        doc.add_paragraph(f"Consistency Ratio (CR): {result['cr']}")
        doc.add_paragraph(f"Random Index (RI): {result['ri']}")
        doc.add_paragraph("\n")

    if aggregate_result:
        doc.add_heading('Aggregate Results', level=2)
        doc.add_paragraph(f"Aggregate Pairwise Comparison Matrix:\n{np.array(aggregate_result['aggregate_matrix'])}")
        doc.add_paragraph(f"Aggregate Weights:\n{aggregate_result['weights']}")
        doc.add_paragraph(f"Aggregate Max Eigenvalue: {aggregate_result['max_eigenvalue']}")
        doc.add_paragraph(f"Aggregate Consistency Index (CI): {aggregate_result['ci']}")
        doc.add_paragraph(f"Aggregate Consistency Ratio (CR): {aggregate_result['cr']}")
        doc.add_paragraph(f"Aggregate Random Index (RI): {aggregate_result['ri']}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    doc_path = f'criteria_analysis_{timestamp}.docx'
    doc.save(doc_path)
    return doc_path

@app.post("/uploadfile/")
async def create_upload_files(files: list[UploadFile] = File(...)):
    try:
        logger.info("Files upload initiated")
        results_list = []
        aggregate_results_list = []

        for file in files:
            content = await file.read()
            content_str = content.decode("utf-8")

            logger.info(f"Processing file content for {file.filename}")
            results, aggregate_result = process_criteria(content_str)
            results_list.extend(results)
            if aggregate_result:
                aggregate_results_list.append(aggregate_result)

        logger.info("Saving analysis results to Word document")
        doc_path = save_to_word(results_list, aggregate_results_list[0] if aggregate_results_list else None)

        response_data = {
            "message": "Files processed successfully",
            "document": doc_path,
            "results": results_list,
            "aggregate_result": aggregate_results_list[0] if aggregate_results_list else None
        }

        logger.info("Files processed successfully")
        return JSONResponse(content=response_data)
    except Exception as e:
        logger.error(f"Error processing files: {e}", exc_info=True)
        return JSONResponse(content={"message": str(e)}, status_code=500)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
